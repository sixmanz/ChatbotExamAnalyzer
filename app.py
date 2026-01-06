# -*- coding: utf-8 -*-
import os
import sys
import subprocess
import json
import re
import io
import time
import pandas as pd
from dotenv import load_dotenv
import streamlit as st
from PyPDF2 import PdfReader 
import google.generativeai as genai
from google.generativeai import types
from google.generativeai.types import GenerationConfig 
from typing import Dict, Any, Tuple
import altair as alt
from shadcn_style import SHADCN_CSS
from datetime import datetime

# Export Libraries
try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

try:
    from fpdf import FPDF
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


# --- 1. การตั้งค่าและโหลด Environment ---
load_dotenv()

# API Keys for all providers
GEMINI_API_KEY = os.getenv('GEMINI_API_KEY', '').strip()
GROQ_API_KEY = os.getenv('GROQ_API_KEY', '').strip()
OPENROUTER_API_KEY = os.getenv('OPENROUTER_API_KEY', '').strip()

# &#128640; Multi-Provider AI Configuration
AI_PROVIDERS = {
    "Gemini (Google)": {
        "models": {
            "Gemini 2.0 Flash (แนะนำ)": "gemini-2.0-flash",
            "Gemini 1.5 Flash (เร็ว)": "gemini-1.5-flash-latest",
            "Gemini 1.5 Pro (แม่นยำ)": "gemini-1.5-pro-latest",
        },
        "api_key": GEMINI_API_KEY,
    },
    "Groq (ฟรี+เร็วมาก)": {
        "models": {
            "Llama 3.3 70B (แนะนำ)": "llama-3.3-70b-versatile",
            "Llama 3.1 8B (เร็ว)": "llama-3.1-8b-instant",
            "Mixtral 8x7B": "mixtral-8x7b-32768",
        },
        "api_key": GROQ_API_KEY,
    },
    "OpenRouter (หลายโมเดลฟรี)": {
        "models": {
            "Llama 3.2 3B (ฟรี)": "meta-llama/llama-3.2-3b-instruct:free",
            "Mistral 7B (ฟรี)": "mistralai/mistral-7b-instruct:free",
            "Gemma 2 9B (ฟรี)": "google/gemma-2-9b-it:free",
        },
        "api_key": OPENROUTER_API_KEY,
    },
}

DEFAULT_PROVIDER = "Gemini (Google)"
DEFAULT_MODEL_NAME = "Gemini 2.0 Flash (แนะนำ)"

# Legacy support - keep AVAILABLE_AI_MODELS for compatibility
AVAILABLE_AI_MODELS = AI_PROVIDERS[DEFAULT_PROVIDER]["models"]

# Check API availability for each provider
GEMINI_AVAILABLE = False
GROQ_AVAILABLE = False
OPENROUTER_AVAILABLE = False

try:
    if GEMINI_API_KEY and len(GEMINI_API_KEY) > 30:
        genai.configure(api_key=GEMINI_API_KEY)
        GEMINI_AVAILABLE = True
except Exception:
    GEMINI_AVAILABLE = False

if GROQ_API_KEY and len(GROQ_API_KEY) > 20:
    GROQ_AVAILABLE = True
    
if OPENROUTER_API_KEY and len(OPENROUTER_API_KEY) > 20:
    OPENROUTER_AVAILABLE = True

def render_user_manual():
    """Show User Manual"""
    with st.expander("&#128218; คู่มือสำหรับมือใหม่ (คลิกอ่าน)", expanded=False):
        st.markdown('''
        **1. เริ่มต้นใช้งาน &#128640;**
        - เลือกไฟล์ข้อสอบ (PDF, DOCX, TXT) จากเครื่อง
        - ระบบจะสกัดโจทย์ออกมาให้ตรวจสอบ
        
        **2. การเลือก AI (สำคัญ) &#129504;**
        - **Gemini**: ฉลาดที่สุด แต่อาจติด Limit (429)
        - **Groq**: เร็วมาก & ฟรี! (แนะนำเมื่อ Gemini เต็ม)
        - *เปลี่ยนได้ที่ "เลือก AI Provider" ด้านล่าง*
        
        **3. ฟีเจอร์พิเศษ &#10024;**
        - **Export**: ดาวน์โหลด Excel/PDF ที่แท็บ Export
        - **สร้างข้อสอบ**: ให้ AI ช่วยคิดโจทย์ใหม่
        - **ประวัติ**: ดูผลย้อนหลังได้ที่ด้านล่างสุด
        
        **4. แก้ปัญหา Error 429 &#9888;**
        - หาก AI ไม่ตอบ (Quota Exceeded)
        - ให้เปลี่ยน Provider เป็น **Groq** หรือ **OpenRouter** ทันที
        ''')


def render_top_navigation():
    """Show Top Navigation Bar (Settings & Manual)"""
    with st.container():
        col1, col2, col3, col4 = st.columns([1, 1, 1.5, 2])
        
        with col1:
            # Language
            st.button(
                t('language_btn'), 
                on_click=toggle_language,
                use_container_width=True,
                key='top_lang_toggle'
            )
            
        with col2:
            # Manual Popover/Expander
            with st.expander("📚 คู่มือ", expanded=False):
                st.markdown(t('tip_1'))
                st.markdown(t('tip_2'))
                st.markdown("---")
                st.markdown(t('quota_warning'))

        with col3:
            # Provider Selector
            provider_options = list(AI_PROVIDERS.keys())
            current_idx = provider_options.index(st.session_state.selected_provider) if st.session_state.selected_provider in provider_options else 0
            
            new_provider = st.selectbox(
                "AI Provider",
                options=provider_options,
                index=current_idx,
                key='top_provider',
                label_visibility="collapsed"
            )
            
            if new_provider != st.session_state.selected_provider:
                st.session_state.selected_provider = new_provider
                # Reset model
                first_model = list(AI_PROVIDERS[new_provider]["models"].keys())[0]
                st.session_state.selected_model = first_model
                st.session_state.analysis_results = None
                st.rerun()

        with col4:
            # Model Selector
            model_options = list(AI_PROVIDERS[st.session_state.selected_provider]["models"].keys())
            current_midx = model_options.index(st.session_state.selected_model) if st.session_state.selected_model in model_options else 0
            
            new_model = st.selectbox(
                "Model",
                options=model_options,
                index=current_midx,
                key='top_model',
                label_visibility="collapsed"
            )
            
            if new_model != st.session_state.selected_model:
                st.session_state.selected_model = new_model
                st.session_state.analysis_results = None

def render_history_sidebar_v2():
    """Show History in Sidebar"""

    history = load_analysis_history()
    
    if not history:
        st.info("ยังไม่มีประวัติ")
        return

    # Show latest first
    for i, entry in enumerate(reversed(history)):
        timestamp = entry.get('timestamp', 'N/A')
        filename = entry.get('filename', 'Unknown')
        
        # Format nice timestamp
        try:
            dt = datetime.fromisoformat(timestamp)
            time_str = dt.strftime("%d/%m %H:%M")
        except:
            time_str = timestamp

        if st.button(f"📂 {filename}", key=f"hist_btn_{i}", use_container_width=True, help=f"วันที่: {time_str}"):
            st.session_state.analysis_results = entry.get('results')
            st.session_state.question_texts = entry.get('question_texts') # Optional restore
            st.success(f"โหลดประวัติ: {filename}")
            st.rerun()


# Initialize session states
if 'analysis_results' not in st.session_state:
    st.session_state.analysis_results = None
if 'last_uploaded_file_name' not in st.session_state:
    st.session_state.last_uploaded_file_name = None
if 'question_texts' not in st.session_state:
    st.session_state.question_texts = None
if 'language' not in st.session_state:
    st.session_state.language = 'th'
if 'custom_prompt' not in st.session_state:
    st.session_state.custom_prompt = ""
if 'selected_provider' not in st.session_state:
    st.session_state.selected_provider = DEFAULT_PROVIDER
if 'selected_model' not in st.session_state:
    st.session_state.selected_model = DEFAULT_MODEL_NAME

# --- Translation Dictionary ---
TRANSLATIONS = {
    'th': {
        # Header
        'app_title': '&#128640; เครื่องมือวิเคราะห์คุณภาพข้อสอบ',
        'app_subtitle': '&#10024; วิเคราะห์ข้อสอบอัตโนมัติด้วย <strong style="color: #667eea;">Gemini AI</strong> ตามหลักสูตรแกนกลางฯ และ <strong style="color: #764ba2;">Bloom\'s Taxonomy</strong>',
        
        # Sidebar
        'sidebar_title': '⚙️ การตั้งค่า & สถานะ AI',
        'ai_connected': '&#9989; เชื่อมต่อ AI สำเร็จ',
        'ai_not_connected': '&#10060; ไม่พบ API Key (GEMINI_API_KEY) - กรุณาตั้งค่าใน `.env`',
        'model_used': '**โมเดลที่ใช้งาน**',
        'batch_analysis': 'Batch Analysis',
        'tips_title': '&#128161; เคล็ดลับ',
        'tip_1': '- ใช้ไฟล์ **PDF** หรือ **TXT**',
        'tip_2': '- ข้อสอบควรมี **เลขข้อ** (เช่น 1., 2.) และ **ตัวเลือก** (เช่น ก., ข.)',
        'api_warning': 'โปรดทราบ: คุณต้องตั้งค่า GEMINI_API_KEY ในไฟล์ .env เพื่อใช้งานส่วนวิเคราะห์',
        
        # Custom Prompt
        'custom_prompt_title': '📝 Custom Prompt (ไม่บังคับ)',
        'custom_prompt_label': 'กรอก Prompt ที่ต้องการใช้แทน Prompt.txt:',
        'custom_prompt_placeholder': 'ปล่อยว่างเพื่อใช้ Prompt จากไฟล์ Prompt.txt...\n\nหรือกรอก Prompt ใหม่ที่นี่ เช่น:\n\nวิเคราะห์ข้อสอบนี้ตามหลัก Bloom\'s Taxonomy และให้คะแนนคุณภาพ...',
        'custom_prompt_active': '&#10024; กำลังใช้ Custom Prompt',
        'custom_prompt_default': '&#128196; กำลังใช้ Prompt จากไฟล์',
        
        # Step 1
        'step1_title': '1️⃣ อัปโหลดไฟล์ข้อสอบ (Batch Analysis)',
        'file_uploader_label': '📁 เลือกไฟล์ข้อสอบ **(.PDF หรือ .TXT)**',
        'reading_file': '⏳ **กำลังอ่านและสกัดข้อสอบ:**',
        'from_file': 'จากไฟล์',
        'extracting': 'กำลังสกัดข้อความ...',
        'no_questions_found': '&#10060; **ไม่พบข้อสอบ** กรุณาตรวจสอบรูปแบบไฟล์ (ไม่มีเลขข้อ/ตัวเลือก หรือรูปแบบซับซ้อนเกินไป)',
        'file_tip': '&#128161; **เคล็ดลับการเตรียมไฟล์:** ไฟล์ควรมี **เลขข้อ** ที่ชัดเจน (เช่น 1., 2., 3.) และมี **ตัวเลือก** (เช่น ก., ข., ค., ง.)',
        'file_read_error': '&#10060; **เกิดข้อผิดพลาดในการอ่านไฟล์:**',
        'extracted_questions': '&#9989; สกัดข้อสอบได้แล้ว **{count} ข้อ** จากไฟล์ `{filename}`',
        
        # Step 2
        'step2_title': '2️⃣ เริ่มต้นวิเคราะห์และสร้างรายงาน &#128640;',
        'start_analysis_btn': '&#128640; **กดที่นี่เพื่อเริ่มการวิเคราะห์ด้วย AI**',
        'api_not_ready': 'Key ไม่พร้อมใช้งาน กรุณาตรวจสอบการตั้งค่า',
        'starting_analysis': '&#128640; **กำลังเริ่มการวิเคราะห์ด้วย AI...**',
        'preparing_analysis': '⏳ กำลังเตรียมวิเคราะห์ข้อสอบ {count} ข้อ โดยใช้ `{model}`',
        'analyzing_question': '&#129302; วิเคราะห์ข้อที่ {num}...',
        'analysis_progress': 'กำลังวิเคราะห์ข้อสอบ {current}/{total} ข้อ...',
        'analysis_complete': '🎉 **การวิเคราะห์เสร็จสมบูรณ์!**',
        
        # Step 3 - Results
        'step3_title': '3️⃣ ผลการวิเคราะห์ชุดข้อสอบ 📝',
        'tab_summary': '📊 สรุปรายงาน & เกณฑ์ Bloom',
        'tab_details': '📝 รายละเอียดรายข้อ',
        'summary_title': '📊 สรุปภาพรวมคุณภาพข้อสอบ',
        'good_questions': '&#9989; ข้อสอบคุณภาพดี',
        'needs_improvement': '&#9888; ข้อสอบต้องปรับปรุง',
        'total_questions': '📝 จำนวนข้อสอบทั้งหมด',
        'analyzed_success': '&#129302; วิเคราะห์สำเร็จ',
        'bloom_criteria_title': '&#128161; เกณฑ์การกระจายระดับความคิด (Bloom)',
        'bloom_low': 'ระดับความคิดต่ำ (จำ/เข้าใจ)',
        'bloom_mid': 'ระดับความคิดกลาง (ใช้/วิเคราะห์)',
        'bloom_high': 'ระดับความคิดสูง (ประเมิน/สร้างสรรค์)',
        'target': 'เป้าหมาย',
        'unidentified_bloom': '**ข้อที่ระบุระดับความคิดไม่ได้:**',
        'bloom_distribution': '📈 การกระจายระดับ Bloom\'s Taxonomy',
        'bloom_table_title': '**ตารางสรุปจำนวนและคุณภาพข้อสอบตามระดับ Bloom**',
        'details_title': '📝 รายละเอียดผลการวิเคราะห์รายข้อ',
        'click_detail': '### 🔎 คลิกดูรายละเอียดการวิเคราะห์ (10 Fields) รายข้อ',
        'question_num': 'ข้อที่',
        'quality': 'คุณภาพข้อสอบ',
        'bloom_level': 'ระดับความคิด',
        'curriculum': 'มาตรฐานหลักสูตร',
        'answer': 'คำตอบ',
        'reasoning': 'เหตุผลโดยย่อ',
        'suggestion': 'ข้อเสนอแนะ',
        'full_question': '**คำถามเต็ม:**',
        'good': '&#9989; คุณภาพดี',
        'improve': '&#10060; ต้องปรับปรุง/ล้มเหลว',
        'difficulty': '⚖️ ความยาก:',
        'correct_answer': '&#9989; คำตอบ:',
        'curriculum_indicator': '**&#128218; ตัวชี้วัดหลักสูตร:**',
        'bloom_reason': '**&#129504; เหตุผลของระดับ Bloom/คุณภาพ:**',
        'answer_analysis_title': 'วิเคราะห์คำตอบและตัวลวง',
        'correct_analysis': '**&#9989; วิเคราะห์คำตอบที่ถูก:**',
        'distractor_analysis': '**&#10060; วิเคราะห์ตัวเลือกลวง (Distractors):**',
        'why_good_distractor': '**&#128161; เหตุผลที่ตัวลวงดี:**',
        'improvement_suggestion': '**🔧 ข้อเสนอแนะในการปรับปรุง:**',
        
        # Quota Warning
        'quota_warning': '&#9888; **ข้อจำกัด Free Tier:** 20 requests/วัน หากเกินโควต้า กรุณารอ 24 ชั่วโมง หรืออัปเกรดแผนการใช้งาน',
        
        # Language
        'language_btn': '🌐 English',
    },
    'en': {
        # Header
        'app_title': '&#128640; Exam Quality Analysis Tool',
        'app_subtitle': '&#10024; Automatic exam analysis with <strong style="color: #667eea;">Gemini AI</strong> based on Core Curriculum and <strong style="color: #764ba2;">Bloom\'s Taxonomy</strong>',
        
        # Sidebar
        'sidebar_title': '⚙️ Settings & AI Status',
        'ai_connected': '&#9989; AI Connected Successfully',
        'ai_not_connected': '&#10060; API Key not found (GEMINI_API_KEY) - Please set in `.env`',
        'model_used': '**Model Used**',
        'batch_analysis': 'Batch Analysis',
        'tips_title': '&#128161; Tips',
        'tip_1': '- Use **PDF** or **TXT** files',
        'tip_2': '- Questions should have **numbers** (e.g., 1., 2.) and **choices** (e.g., A., B.)',
        'api_warning': 'Note: You must set GEMINI_API_KEY in .env file to use the analysis feature',
        
        # Custom Prompt
        'custom_prompt_title': '📝 Custom Prompt (Optional)',
        'custom_prompt_label': 'Enter custom prompt to use instead of Prompt.txt:',
        'custom_prompt_placeholder': 'Leave empty to use default Prompt.txt...\n\nOr enter your custom prompt here, e.g.:\n\nAnalyze this exam question according to Bloom\'s Taxonomy and rate its quality...',
        'custom_prompt_active': '&#10024; Using Custom Prompt',
        'custom_prompt_default': '&#128196; Using Default Prompt File',
        
        # Step 1
        'step1_title': '1️⃣ Upload Exam File (Batch Analysis)',
        'file_uploader_label': '📁 Select exam file **(.PDF or .TXT)**',
        'reading_file': '⏳ **Reading and extracting questions:**',
        'from_file': 'from file',
        'extracting': 'Extracting text...',
        'no_questions_found': '&#10060; **No questions found** Please check the file format (no question numbers/choices or format too complex)',
        'file_tip': '&#128161; **File preparation tip:** File should have clear **question numbers** (e.g., 1., 2., 3.) and **choices** (e.g., A., B., C., D.)',
        'file_read_error': '&#10060; **Error reading file:**',
        'extracted_questions': '&#9989; Extracted **{count} questions** from file `{filename}`',
        
        # Step 2
        'step2_title': '2️⃣ Start Analysis & Generate Report &#128640;',
        'start_analysis_btn': '&#128640; **Click here to start AI analysis**',
        'api_not_ready': 'API Key not ready. Please check settings',
        'starting_analysis': '&#128640; **Starting AI analysis...**',
        'preparing_analysis': '⏳ Preparing to analyze {count} questions using `{model}`',
        'analyzing_question': '&#129302; Analyzing question {num}...',
        'analysis_progress': 'Analyzing question {current}/{total}...',
        'analysis_complete': '🎉 **Analysis Complete!**',
        
        # Step 3 - Results
        'step3_title': '3️⃣ Exam Analysis Results 📝',
        'tab_summary': '📊 Summary & Bloom Criteria',
        'tab_details': '📝 Question Details',
        'summary_title': '📊 Overall Exam Quality Summary',
        'good_questions': '&#9989; Good Quality Questions',
        'needs_improvement': '&#9888; Needs Improvement',
        'total_questions': '📝 Total Questions',
        'analyzed_success': '&#129302; Successfully Analyzed',
        'bloom_criteria_title': '&#128161; Bloom\'s Taxonomy Distribution Criteria',
        'bloom_low': 'Lower Order (Remember/Understand)',
        'bloom_mid': 'Middle Order (Apply/Analyze)',
        'bloom_high': 'Higher Order (Evaluate/Create)',
        'target': 'Target',
        'unidentified_bloom': '**Questions with unidentified Bloom level:**',
        'bloom_distribution': '📈 Bloom\'s Taxonomy Distribution',
        'bloom_table_title': '**Summary Table: Questions by Bloom Level**',
        'details_title': '📝 Detailed Analysis per Question',
        'click_detail': '### 🔎 Click to view detailed analysis (10 Fields) per question',
        'question_num': 'Q#',
        'quality': 'Quality',
        'bloom_level': 'Bloom Level',
        'curriculum': 'Curriculum Standard',
        'answer': 'Answer',
        'reasoning': 'Brief Reasoning',
        'suggestion': 'Suggestion',
        'full_question': '**Full Question:**',
        'good': '&#9989; Good',
        'improve': '&#10060; Needs Improvement/Failed',
        'difficulty': '⚖️ Difficulty:',
        'correct_answer': '&#9989; Answer:',
        'curriculum_indicator': '**&#128218; Curriculum Indicator:**',
        'bloom_reason': '**&#129504; Bloom Level/Quality Reasoning:**',
        'answer_analysis_title': 'Answer & Distractor Analysis',
        'correct_analysis': '**&#9989; Correct Answer Analysis:**',
        'distractor_analysis': '**&#10060; Distractor Analysis:**',
        'why_good_distractor': '**&#128161; Why Good Distractors:**',
        'improvement_suggestion': '**🔧 Improvement Suggestion:**',
        
        # Quota Warning
        'quota_warning': '&#9888; **Free Tier Limit:** 20 requests/day. If exceeded, please wait 24 hours or upgrade your plan.',
        
        # Language
        'language_btn': '🌐 ภาษาไทย',
    }
}

def t(key):
    """Get translation for current language"""
    lang = st.session_state.get('language', 'th')
    return TRANSLATIONS.get(lang, TRANSLATIONS['th']).get(key, key)


# --- 2. การโหลด Prompt Template ---
def load_prompts() -> Tuple[str, str]:
    """อ่านเนื้อหา Prompt Template จากไฟล์ 'Prompt.txt' และแยกเป็น System/Chat"""
    prompt_path = os.path.join(os.path.dirname(__file__), 'Prompt.txt')
    try:
        with open(prompt_path, 'r', encoding='utf-8') as f:
            full_content = f.read()
            
            # 1. System Instruction
            system_match = full_content.split("# --- END_SYSTEM_INSTRUCTION_ANALYSIS_MODE ---")[0]
            SYSTEM_INSTRUCTION_PROMPT = system_match.strip() if system_match else "You are a test expert for Thai curriculum. Analyze the question and return a JSON object."
            
            # 2. Few-Shot Template 
            chat_match = full_content.split("# --- CHAT_PROMPT ---")
            
            if len(chat_match) > 1:
                # แยกส่วน Prompt Template ทั้งหมด
                template_content = chat_match[1].split("# --- CHAT_PROMPT_END ---")[0].strip()
                # และเพิ่ม {user_query} เข้าไปในส่วนท้าย
                FEW_SHOT_PROMPT_TEMPLATE = template_content + "\n{user_query}"
            else:
                FEW_SHOT_PROMPT_TEMPLATE = "Analyze the following question/text and return the JSON object: {user_query}"

            return SYSTEM_INSTRUCTION_PROMPT, FEW_SHOT_PROMPT_TEMPLATE
    except FileNotFoundError:
        return (
            "# Error: Prompt file not found. Fallback to default instruction.", 
            "Analyze the following question/text and return the JSON object: {user_query}"
        )

SYSTEM_INSTRUCTION_PROMPT, FEW_SHOT_PROMPT_TEMPLATE = load_prompts()


# --- 3. Helper Functions ---

BLOOM_COLORS = {
    "REMEMBER": "#6495ED",  
    "UNDERSTAND": "#40E0D0",
    "APPLY": "#50C878",     
    "ANALYZE": "#8FBC8F",   
    "EVALUATE": "#ADFF2F",  
    "CREATE": "#FFD700",    
    "ไม่ระบุ": "#A9A9A9",
    "ไม่สามารถระบุได้": "#A9A9A9"
}

def get_bloom_color(level):
    # Returns Hex color for Bloom's Taxonomy level
    if not level:
        return BLOOM_COLORS['ไม่ระบุ']
    
    level_upper = level.strip().upper()
    for key, color in BLOOM_COLORS.items():
        if key in level_upper:
            return color
    return BLOOM_COLORS['ไม่ระบุ'] # Fallback

def get_text_color_for_bloom(level):
    """กำหนดสีตัวอักษรให้ตัดกับสีพื้นหลัง (เพื่อให้มองเห็นได้ชัดเจน)"""
    level_upper = level.strip().upper()
    if level_upper in ['EVALUATE', 'CREATE']: # สีเหลือง/ทอง/สว่าง ควรใช้ตัวอักษรสีดำ
        return 'black' 
    return 'white' # สีเข้มอื่นๆ ใช้ตัวอักษรสีขาว


def extract_text_from_pdf(pdf_reader):
    """สกัดข้อความจากวัตถุ PdfReader"""
    text = ""
    for page in pdf_reader.pages:
        try:
            page_text = page.get_text() if hasattr(page, 'get_text') else page.extract_text()
            text += (page_text or "") + "\n"
        except Exception:
            text += "\n"
    return text.strip()


def clean_and_normalize(text):
    """ทำความสะอาดข้อความและแปลงเลขไทยเป็นเลขอารบิก"""
    if not text: return ""
    # 1. แปลงเลขไทย
    thai_digits = "๐๑๒๓๔๕๖๗๘๙"
    for i, digit in enumerate(thai_digits):
        text = text.replace(digit, str(i))
    
    # 2. ทำความสะอาดตัวอักษรพิเศษและช่องว่าง
    text = re.sub(r'[ \t]+', ' ', text) 
    text = text.replace('\r', '')
    
    # 3. ทำให้ตัวเลือกติดกับจุด (เช่น ก . -> ก.)
    text = re.sub(r'([ก-งA-D])\s*\.', r'\1.', text)
    
    # 4. ทำให้เลขข้อติดกับจุด (เช่น 1 . -> 1.)
    text = re.sub(r'(\d+)\s*\.', r'\1.', text)
    
    # NEW: Insert newline before potential question start if preceded by punctuation or space
    # Matches: "text. 2. text" -> "text.\n2. text"
    # Matches: "text (2) text" -> "text\n(2) text"
    text = re.sub(r'(\s+)(\(?\d+[\.\)])\s', r'\n\2 ', text)
    
    # 5. ลบบรรทัดว่างที่ติดกันหลายบรรทัด
    text = re.sub(r'\n{2,}', '\n', text)
    
    lines = [line.strip() for line in text.split('\n')]
    return '\n'.join(lines)


def extract_questions_with_ai(raw_text):
    """
    Fallback: ให้ AI ช่วยแยกข้อสอบเมื่อ Regex เอาไม่อยู่
    """
    if not GEMINI_AVAILABLE:
        return []

    try:
        model = genai.GenerativeModel("gemini-1.5-flash-latest") # Use Flash for speed/cost
        
        prompt = f"""
        You are an expert exam parser. 
        Please extract all exam questions from the following text and return them as a JSON list of strings.
        
        Rules:
        1. Capture the full question text including the question number and all options (e.g. "1. Question... A. Opt...").
        2. Do not change the original text, just split it correctly.
        3. If there are no clear questions, return an empty list.
        4. Return ONLY raw JSON Array.

        Text to parse:
        {raw_text[:20000]} 
        """
        # Limit text to 20k chars to avoid token limits on fallback
        
        response = model.generate_content(prompt, generation_config={"response_mime_type": "application/json"})
        questions = json.loads(response.text)
        
        if isinstance(questions, list):
            return [str(q).strip() for q in questions]
        else:
            return []
            
    except Exception as e:
        print(f"AI Extraction Failed: {e}")
        return []

def extract_questions(raw_text):
    """
    สกัดข้อสอบเป็นรายข้อ (ปรับปรุงให้รองรับหลายรูปแบบ: 1., 1), (1), ข้อ 1, ข้อที่ 1)
    """
    # 1. ทำความสะอาดข้อความทั้งหมด
    text = re.split(r"={10,}\s*เฉลย\s*={10,}", raw_text, flags=re.DOTALL | re.IGNORECASE)[0]
    cleaned_text = clean_and_normalize(text)
    
    # 2. Regex สำหรับจับเลขข้อ
    # รองรับ: "1.", "1)", "(1)", "ข้อ 1", "ข้อที่ 1", "ข้อ ๑", (และแบบมี space)
    question_pattern = r'(?:^|\n)\s*((?:ข้อ\s*\d+|ข้อที่\s*\d+|\d+\.|(?:\(?\d+\)))[\.\s]+)'
    
    chunks = re.split(question_pattern, cleaned_text)
    
    questions = []
    
    # Skip preamble
    start_idx = 0
    if len(chunks) > 0 and not chunks[0].strip():
        start_idx = 1
    elif len(chunks) > 0 and not re.match(r'(?:ข้อ\s*\d+|ข้อที่\s*\d+|\d+\.|(?:\(?\d+\)))', chunks[0].strip()):
        start_idx = 1 # Skip likely header

    for i in range(start_idx, len(chunks), 2):
        if i+1 < len(chunks):
            delim = chunks[i]
            content = chunks[i+1]
            full_q = delim + content
            questions.append(full_q.strip())

    # 3. การตรวจสอบความถูกต้อง (Validation แบบยืดหยุ่น)
    valid_questions = []
    for q in questions:
        q = q.strip()
        if len(q) < 5: continue 
        
        has_std_options = len(re.findall(r'[ก-งA-D]\.', q)) >= 2
        
        if has_std_options:
            q_formatted = re.sub(r'(\s+)([ก-งA-D]\.)', r'\n\2', q)
            valid_questions.append(q_formatted)
        else:
            if len(q) > 10: 
                valid_questions.append(q)
    
    # --- FALLBACK TO AI ---
    # ถ้าหาไม่เจอเลย หรือเจอน้อยผิดปกติเมื่อเทียบกับความยาวข้อความ (เช่น ข้อความยาว 5000 ตัวอักษร แต่เจอ 0 ข้อ)
    # หรือถ้าเจอ < 3 ข้อ แต่ข้อความยาวมาก ให้ลองใช้ AI ช่วย
    is_suspiciously_low = len(valid_questions) == 0 or (len(valid_questions) < 3 and len(raw_text) > 500)
    
    if is_suspiciously_low and GEMINI_AVAILABLE:
        st.toast("&#9888; รูปแบบไฟล์ซับซ้อน กำลังใช้ AI ช่วยแกะข้อสอบ...", icon="&#129302;")
        ai_extracted = extract_questions_with_ai(raw_text)
        if len(ai_extracted) > len(valid_questions):
             return ai_extracted

    return valid_questions


def analyze_with_gemini(question_text, question_id=1):
    """เรียกใช้ Gemini API เพื่อวิเคราะห์ข้อสอบ (10 Fields Logic)"""
    if not GEMINI_AVAILABLE:
        return {
            "bloom_level": "ไม่ระบุ", "reasoning": "ไม่พบ API Key",
            "difficulty": "ไม่ระบุ", "curriculum_standard": "ไม่ระบุ",
            "correct_option": "ไม่ระบุ", "correct_option_analysis": "ไม่มีการเชื่อมต่อ AI",
            "distractor_analysis": "ไม่มีการเชื่อมต่อ AI", "why_good_distractor": "ไม่มีการเชื่อมต่อ AI",
            "is_good_question": False, "improvement_suggestion": "ไม่สามารถวิเคราะห์ได้: ไม่พบ API Key"
        } 

    # Check for custom prompt
    custom_prompt = st.session_state.get('custom_prompt', '').strip()
    
    if custom_prompt:
        # Use custom prompt as system instruction
        system_instruction = custom_prompt
        prompt_template = "{user_query}"  # Simple template for custom prompt
    else:
        # Use default prompts from Prompt.txt
        system_instruction = SYSTEM_INSTRUCTION_PROMPT
        prompt_template = FEW_SHOT_PROMPT_TEMPLATE

    # ใช้โมเดลที่ผู้ใช้เลือก
    selected_model_name = st.session_state.get('selected_model', DEFAULT_MODEL_NAME)
    model_id = AVAILABLE_AI_MODELS.get(selected_model_name, "gemini-2.0-flash")
    
    model = genai.GenerativeModel(
        model_id, 
        system_instruction=system_instruction
    )
    
    question_text_formatted = f"คำถามข้อที่ {question_id}:\n{question_text}"
    
    # การจัดรูปแบบ Prompt ที่ถูกต้อง
    full_prompt = prompt_template.format(user_query=question_text_formatted) 
    
    # กำหนด JSON Schema (10 Fields)
    json_schema = {
        "type": "object",
        "properties": {
            "bloom_level": {"type": "string"},
            "reasoning": {"type": "string"},
            "difficulty": {"type": "string"},
            "curriculum_standard": {"type": "string"},
            "correct_option": {"type": "string"},
            "correct_option_analysis": {"type": "string"},
            "distractor_analysis": {"type": "string"},
            "why_good_distractor": {"type": "string"},
            "is_good_question": {"type": "boolean"},
            "improvement_suggestion": {"type": "string"}
        },
        "required": [
            "bloom_level", "reasoning", "difficulty", "curriculum_standard",
            "correct_option", "correct_option_analysis", "distractor_analysis",
            "why_good_distractor", "is_good_question", "improvement_suggestion"
        ]
    }

    # กำหนด Config (บังคับ JSON และลด Temperature)
    config = GenerationConfig( 
        response_mime_type="application/json", 
        response_schema=json_schema, 
        temperature=0.2  # เพิ่มเล็กน้อยเพื่อให้ได้ผลลัพธ์ที่หลากหลายขึ้น
    )
    
    last_error_message = ""
    max_retries = 5  # เพิ่มจำนวน retry

    for attempt in range(max_retries):
        # Exponential backoff with jitter
        if attempt > 0:
            base_delay = min(60, (2 ** attempt) + (attempt * 2))  # 2, 6, 12, 22, 38 seconds
            time.sleep(base_delay)
            
        try:
            response = model.generate_content(
                full_prompt, 
                generation_config=config, 
            )
            raw_text = response.text.strip()
        
            # Hardened JSON Cleaning/Extraction
            cleaned_json = re.sub(r'^```(?:json)?\s*|```$', '', raw_text, flags=re.MULTILINE | re.DOTALL).strip()
            start_brace = cleaned_json.find('{')
            end_brace = cleaned_json.rfind('}')
            
            if start_brace != -1 and end_brace != -1 and end_brace > start_brace:
                cleaned_json = cleaned_json[start_brace:end_brace+1].strip()
            else:
                raise ValueError("Could not find valid JSON structure (missing { or }).") 
            
            # โหลด JSON
            analysis = json.loads(cleaned_json)
            
            # Data sanitation via shared function
            final_analysis = sanitize_analysis(analysis)
            return final_analysis

        except (json.JSONDecodeError, ValueError, KeyError) as e:
            last_error_message = f"ข้อผิดพลาดในการประมวลผล JSON/Key: {type(e).__name__}: {str(e)}"
            if attempt < max_retries - 1:
                continue  # ลองใหม่ (Retry)
            
        except Exception as e:
            error_str = str(e).lower()
            
            # Handle Rate Limit / Quota Exceeded (429 Error)
            is_rate_limit = any([
                "429" in str(e),
                "quota" in error_str,
                "resourceexhausted" in error_str,
                "rate" in error_str and "limit" in error_str,
                "too many requests" in error_str
            ])
            
            if is_rate_limit:
                # Calculate progressive delay
                if attempt < max_retries - 1:
                    retry_delay = min(120, 15 * (attempt + 1))  # 15, 30, 45, 60, 75 seconds
                    last_error_message = f"⏳ Rate Limit: รอ {retry_delay} วินาที แล้วลองใหม่... (ครั้งที่ {attempt + 1}/{max_retries})"
                    time.sleep(retry_delay)
                    continue
                else:
                    last_error_message = f"&#10060; Quota Exceeded: คุณใช้โควต้า API ครบแล้ว กรุณารอสักครู่หรือสร้าง API Key ใหม่"
                    break
            else:
                last_error_message = f"ข้อผิดพลาด: {type(e).__name__}: {str(e)}"
                if attempt < max_retries - 1:
                    continue

    # Fallback สุดท้าย: หาก AI วิเคราะห์ไม่ได้เลย
    return {
        "bloom_level": "ไม่สามารถระบุได้", "reasoning": "AI วิเคราะห์ล้มเหลว",
        "difficulty": "ไม่สามารถประเมินได้", "curriculum_standard": "ไม่สามารถระบุได้",
        "correct_option": "ไม่ระบุ", "correct_option_analysis": "ไม่ระบุ",
        "distractor_analysis": "ไม่ระบุ", "why_good_distractor": "ไม่ระบุ",
        "is_good_question": False, 
        "improvement_suggestion": f"**เกิดข้อผิดพลาดในการวิเคราะห์**: {last_error_message}"
    }


def analyze_with_groq(question_text, question_id=1):
    """วิเคราะห์ข้อสอบผ่าน Groq API (Llama, Mixtral) [Robust]"""
    from groq import Groq
    
    if not GROQ_AVAILABLE:
        return create_error_response("ไม่พบ GROQ_API_KEY")
    
    client = Groq(api_key=GROQ_API_KEY)
    
    # Get selected model
    selected_model = st.session_state.get('selected_model', 'Llama 3.3 70B (แนะนำ)')
    model_id = AI_PROVIDERS["Groq (ฟรี+เร็วมาก)"]["models"].get(selected_model, "llama-3.3-70b-versatile")
    
    # Build prompt
    custom_prompt = st.session_state.get('custom_prompt', '').strip()
    system_prompt = custom_prompt if custom_prompt else SYSTEM_INSTRUCTION_PROMPT
    
    user_message = f"""คำถามข้อที่ {question_id}:
{question_text}

วิเคราะห์และตอบเป็น JSON ที่มี keys: bloom_level, reasoning, difficulty, curriculum_standard, correct_option, correct_option_analysis, distractor_analysis, why_good_distractor, is_good_question (boolean), improvement_suggestion"""
    
    max_retries = 3
    last_error = ""
    
    for attempt in range(max_retries):
        if attempt > 0:
            time.sleep(attempt * 2)
            
        try:
            response = client.chat.completions.create(
                model=model_id,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_message}
                ],
                temperature=0.2,
                max_tokens=2000,
                response_format={"type": "json_object"}
            )
            
            raw_text = response.choices[0].message.content
            analysis = json.loads(raw_text)
            return sanitize_analysis(analysis)
            
        except Exception as e:
            error_str = str(e).lower()
            if "429" in error_str or "rate limit" in error_str:
                last_error = f"Rate Limit (รอ {attempt*2}s)"
                time.sleep(5) # Extra wait for rate limit
            else:
                last_error = str(e)
            
    return create_error_response(f"Groq Error (Max Retries): {last_error}")



def analyze_with_openrouter(question_text, question_id=1):
    """วิเคราะห์ข้อสอบผ่าน OpenRouter API [Robust]"""
    import openai
    
    if not OPENROUTER_AVAILABLE:
        return create_error_response("ไม่พบ OPENROUTER_API_KEY")
    
    client = openai.OpenAI(
        base_url="https://openrouter.ai/api/v1",
        api_key=OPENROUTER_API_KEY
    )
    
    # Get selected model
    selected_model = st.session_state.get('selected_model', 'Llama 3.2 3B (ฟรี)')
    model_id = AI_PROVIDERS["OpenRouter (หลายโมเดลฟรี)"]["models"].get(selected_model, "meta-llama/llama-3.2-3b-instruct:free")
    
    # Build prompt  
    custom_prompt = st.session_state.get('custom_prompt', '').strip()
    system_prompt = custom_prompt if custom_prompt else SYSTEM_INSTRUCTION_PROMPT
    
    user_message = f"""คำถามข้อที่ {question_id}:
{question_text}

วิเคราะห์และตอบเป็น JSON ที่มี keys: bloom_level, reasoning, difficulty, curriculum_standard, correct_option, correct_option_analysis, distractor_analysis, why_good_distractor, is_good_question (boolean), improvement_suggestion"""
    
    max_retries = 3
    last_error = ""
    
    for attempt in range(max_retries):
        if attempt > 0:
            time.sleep(attempt * 2)
            
        try:
            response = client.chat.completions.create(
                model=model_id,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_message}
                ],
                temperature=0.2,
                max_tokens=2000
            )
            
            raw_text = response.choices[0].message.content
            # Clean JSON from markdown code blocks
            cleaned = re.sub(r'^```(?:json)?\s*|```$', '', raw_text, flags=re.MULTILINE | re.DOTALL).strip()
            start_brace = cleaned.find('{')
            end_brace = cleaned.rfind('}')
            if start_brace != -1 and end_brace > start_brace:
                cleaned = cleaned[start_brace:end_brace+1]
            analysis = json.loads(cleaned)
            return sanitize_analysis(analysis)
            
        except Exception as e:
            last_error = str(e)
            if "429" in str(e):
                 time.sleep(5)

    return create_error_response(f"OpenRouter Error: {last_error}")



def create_error_response(error_message):
    """สร้าง response เมื่อเกิดข้อผิดพลาด"""
    return {
        "bloom_level": "ไม่สามารถระบุได้", "reasoning": "AI วิเคราะห์ล้มเหลว",
        "difficulty": "ไม่สามารถประเมินได้", "curriculum_standard": "ไม่สามารถระบุได้",
        "correct_option": "ไม่ระบุ", "correct_option_analysis": "ไม่ระบุ",
        "distractor_analysis": "ไม่ระบุ", "why_good_distractor": "ไม่ระบุ",
        "is_good_question": False, 
        "improvement_suggestion": f"**เกิดข้อผิดพลาด**: {error_message}"
    }


def sanitize_analysis(analysis):
    """ทำความสะอาดและตรวจสอบผลลัพธ์จาก AI (Robust)"""
    required_keys = [
        "bloom_level", "reasoning", "difficulty", "curriculum_standard",
        "correct_option", "correct_option_analysis", "distractor_analysis",
        "why_good_distractor", "is_good_question", "improvement_suggestion"
    ]
    
    result = {}
    missing_keys = []
    
    for key in required_keys:
        val = analysis.get(key)
        
        # Robust Boolean Conversion
        if key == "is_good_question":
            if isinstance(val, bool):
                result[key] = val
            elif isinstance(val, str):
                result[key] = val.strip().lower() in ['true', 'yes', '1', 'correct', 'จริง', 'ใช่']
            else:
                result[key] = False
        else:
            # Robust String Conversion
            result[key] = str(val).strip() if val not in [None, "", "null"] else "ไม่ระบุ"
            
        if key not in analysis:
            missing_keys.append(key)

    # Specific Default Values for missing keys
    if "improvement_suggestion" not in result or result["improvement_suggestion"] == "ไม่ระบุ":
        result["improvement_suggestion"] = "ไม่มีข้อเสนอแนะเพิ่มเติม"

    return result



def analyze_question(question_text, question_id=1):
    """Wrapper function - เลือก provider ตามที่ผู้ใช้เลือก"""
    provider = st.session_state.get('selected_provider', DEFAULT_PROVIDER)
    
    if provider == "Gemini (Google)":
        return analyze_with_gemini(question_text, question_id)
    elif provider == "Groq (ฟรี+เร็วมาก)":
        return analyze_with_groq(question_text, question_id)
    elif provider == "OpenRouter (หลายโมเดลฟรี)":
        return analyze_with_openrouter(question_text, question_id)
    else:
        return analyze_with_gemini(question_text, question_id)


# ===== EXPORT FUNCTIONS =====
def export_to_excel(analysis_results, filename="exam_analysis.xlsx"):
    """Export ผลวิเคราะห์เป็น Excel"""
    if not EXCEL_AVAILABLE:
        return None
    
    wb = Workbook()
    ws = wb.active
    ws.title = "ผลวิเคราะห์ข้อสอบ"
    
    # Header styling
    header_fill = PatternFill(start_color="18181B", end_color="18181B", fill_type="solid")
    header_font = Font(color="FFFFFF", bold=True)
    
    headers = ["ข้อที่", "ระดับ Bloom", "ความยาก", "คุณภาพ", "มาตรฐาน", "คำตอบ", "ข้อเสนอแนะ"]
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center")
    
    # Data rows
    for idx, item in enumerate(analysis_results, 1):
        ws.cell(row=idx+1, column=1, value=idx)
        ws.cell(row=idx+1, column=2, value=item.get('bloom_level', 'N/A'))
        ws.cell(row=idx+1, column=3, value=item.get('difficulty', 'N/A'))
        ws.cell(row=idx+1, column=4, value="ดี" if item.get('is_good_question') else "ต้องปรับปรุง")
        ws.cell(row=idx+1, column=5, value=item.get('curriculum_standard', 'N/A'))
        ws.cell(row=idx+1, column=6, value=item.get('correct_option', 'N/A'))
        ws.cell(row=idx+1, column=7, value=item.get('improvement_suggestion', 'N/A'))
    
    # Auto-adjust column widths
    for col in ws.columns:
        max_length = max(len(str(cell.value or "")) for cell in col)
        ws.column_dimensions[col[0].column_letter].width = min(max_length + 2, 50)
    
    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    return output


def extract_text_from_docx(file):
    """สกัดข้อความจากไฟล์ DOCX (อ่านพารากราฟและตาราง)"""
    if not DOCX_AVAILABLE:
        return None
    try:
        doc = Document(file)
        full_text = []
        
        # อ่าน paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
                
        # อ่าน tables (มักใช้ในข้อสอบ)
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    full_text.append(" ".join(row_text))
                    
        return "\n".join(full_text)
    except Exception as e:
        return None



def generate_exam_with_ai(subject, bloom_level, num_questions, difficulty="ปานกลาง"):
    """สร้างข้อสอบใหม่ด้วย AI"""
    provider = st.session_state.get('selected_provider', DEFAULT_PROVIDER)
    
    prompt = f"""สร้างข้อสอบปรนัย 4 ตัวเลือก จำนวน {num_questions} ข้อ
วิชา: {subject}
Level: {bloom_level}
ระดับความยาก: {difficulty}

สำหรับแต่ละข้อ ให้มี:
1. คำถามที่ชัดเจน
2. ตัวเลือก ก. ข. ค. ง.
3. เฉลย
4. คำอธิบายคำตอบ

ตอบเป็น JSON array ที่มี keys: question, options (array), answer, explanation"""
    
    try:
        if provider == "Gemini (Google)" and GEMINI_AVAILABLE:
            model = genai.GenerativeModel("gemini-2.0-flash")
            response = model.generate_content(prompt)
            raw_text = response.text
        elif provider == "Groq (ฟรี+เร็วมาก)" and GROQ_AVAILABLE:
            from groq import Groq
            client = Groq(api_key=GROQ_API_KEY)
            response = client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.7
            )
            raw_text = response.choices[0].message.content
        elif provider == "OpenRouter (หลายโมเดลฟรี)" and OPENROUTER_AVAILABLE:
            import openai
            client = openai.OpenAI(base_url="https://openrouter.ai/api/v1", api_key=OPENROUTER_API_KEY)
            response = client.chat.completions.create(
                model="meta-llama/llama-3.2-3b-instruct:free",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.7
            )
            raw_text = response.choices[0].message.content
        else:
            return None, "ไม่มี API Key ที่พร้อมใช้งาน"
        
        # Parse JSON from response
        cleaned = re.sub(r'^```(?:json)?\s*|```$', '', raw_text, flags=re.MULTILINE | re.DOTALL).strip()
        start = cleaned.find('[')
        end = cleaned.rfind(']') + 1
        if start != -1 and end > start:
            exams = json.loads(cleaned[start:end])
            return exams, None
        return None, "ไม่สามารถ parse JSON ได้"
    except Exception as e:
        return None, str(e)


def improve_question_with_ai(question_text, suggestion):
    """ปรับปรุงข้อสอบตามคำแนะนำ AI"""
    provider = st.session_state.get('selected_provider', DEFAULT_PROVIDER)
    
    prompt = f"""ข้อสอบเดิม:
{question_text}

ข้อเสนอแนะในการปรับปรุง:
{suggestion}

กรุณาเขียนข้อสอบใหม่ที่ปรับปรุงตามข้อเสนอแนะ โดยยังคงเนื้อหาหลักไว้ แต่แก้ไขจุดบกพร่อง
ตอบเฉพาะข้อสอบที่ปรับปรุงแล้วเท่านั้น ในรูปแบบ:
- คำถาม
- ตัวเลือก ก. ข. ค. ง.
- (เฉลย: ตัวเลือกที่ถูกต้อง)"""
    
    try:
        if provider == "Gemini (Google)" and GEMINI_AVAILABLE:
            model = genai.GenerativeModel("gemini-2.0-flash")
            response = model.generate_content(prompt)
            return response.text, None
        elif provider == "Groq (ฟรี+เร็วมาก)" and GROQ_AVAILABLE:
            from groq import Groq
            client = Groq(api_key=GROQ_API_KEY)
            response = client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.5
            )
            return response.choices[0].message.content, None
        elif provider == "OpenRouter (หลายโมเดลฟรี)" and OPENROUTER_AVAILABLE:
            import openai
            client = openai.OpenAI(base_url="https://openrouter.ai/api/v1", api_key=OPENROUTER_API_KEY)
            response = client.chat.completions.create(
                model="meta-llama/llama-3.2-3b-instruct:free",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.5
            )
            return response.choices[0].message.content, None
        return None, "ไม่มี API Key ที่พร้อมใช้งาน"
    except Exception as e:
        return None, str(e)


def save_analysis_history(filename, results, summary):
    """บันทึกประวัติการวิเคราะห์"""
    history_file = "analysis_history.json"
    history = []
    
    if os.path.exists(history_file):
        try:
            with open(history_file, 'r', encoding='utf-8') as f:
                history = json.load(f)
        except:
            history = []
    
    entry = {
        "timestamp": datetime.now().isoformat(),
        "filename": filename,
        "total_questions": len(results),
        "good_questions": sum(1 for r in results if r.get('is_good_question')),
        "summary": summary
    }
    history.insert(0, entry)  # Add to beginning
    history = history[:20]  # Keep only last 20
    
    with open(history_file, 'w', encoding='utf-8') as f:
        json.dump(history, f, ensure_ascii=False, indent=2)


def load_analysis_history():
    """โหลดประวัติการวิเคราะห์"""
    history_file = "analysis_history.json"
    if os.path.exists(history_file):
        try:
            with open(history_file, 'r', encoding='utf-8') as f:
                return json.load(f)
        except:
            return []
    return []


def extract_text_from_docx(file):
    """สกัดข้อความจากไฟล์ DOCX"""
    if not DOCX_AVAILABLE:
        return None
    try:
        doc = DocxDocument(file)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return '\n'.join(full_text)
    except Exception as e:
        return None


def check_bloom_criteria(analysis_results):
    """ตรวจสอบว่าชุดข้อสอบผ่านเกณฑ์การกระจายระดับ Bloom หรือไม่."""
    total = len(analysis_results)
    if total == 0: 
        return {"pass": False, "reason": "ไม่มีข้อมูลข้อสอบ", "percentages": {}, "raw_counts": {}, "valid_total": 0} 

    # Standardized Level Names for grouping
    LEVEL_GROUPS = {
        "Remember": 0, "Understand": 0, "Apply": 0, 
        "Analyze": 0, "Evaluate": 0, "Create": 0, "Unknown": 0
    }
    
    for item in analysis_results:
        level = str(item.get("bloom_level", "")).strip().lower()
        if "remember" in level: LEVEL_GROUPS["Remember"] += 1
        elif "understand" in level: LEVEL_GROUPS["Understand"] += 1
        elif "apply" in level: LEVEL_GROUPS["Apply"] += 1
        elif "analyze" in level: LEVEL_GROUPS["Analyze"] += 1
        elif "evaluate" in level: LEVEL_GROUPS["Evaluate"] += 1
        elif "create" in level: LEVEL_GROUPS["Create"] += 1
        else: LEVEL_GROUPS["Unknown"] += 1

    apply_analyze_count = LEVEL_GROUPS["Apply"] + LEVEL_GROUPS["Analyze"]
    remember_understand_count = LEVEL_GROUPS["Remember"] + LEVEL_GROUPS["Understand"]
    evaluate_create_count = LEVEL_GROUPS["Evaluate"] + LEVEL_GROUPS["Create"]
    valid_total = total - LEVEL_GROUPS["Unknown"] # ข้อที่สามารถระบุระดับ Bloom ได้

    if valid_total == 0:
        return {"pass": False, "reason": "ไม่มีข้อสอบที่สามารถวิเคราะห์ระดับความคิดได้เลย", "percentages": {}, "raw_counts": LEVEL_GROUPS, "valid_total": 0} 
        
    percent = {
        "Remember/Understand": round((remember_understand_count / valid_total) * 100, 1) if valid_total > 0 else 0,
        "Apply/Analyze": round((apply_analyze_count / valid_total) * 100, 1) if valid_total > 0 else 0,
        "Evaluate/Create": round((evaluate_create_count / valid_total) * 100, 1) if valid_total > 0 else 0
    }

    # เกณฑ์ (ตัวอย่าง: ต่ำ ≤ 40%, กลาง ≥ 50%, สูง ≥ 10%)
    pass_r_u = percent["Remember/Understand"] <= 40
    pass_a_a = percent["Apply/Analyze"] >= 50
    pass_e_c = percent["Evaluate/Create"] >= 10
    passed = pass_r_u and pass_a_a and pass_e_c

    reason = "ชุดข้อสอบ **ผ่าน** เกณฑ์การกระจายระดับความคิด (Bloom)" if passed else "ชุดข้อสอบ **ไม่ผ่าน** เกณฑ์การกระจายระดับความคิด (Bloom)"

    return {
        "pass": passed,
        "reason": reason,
        "percentages": percent,
        "raw_counts": LEVEL_GROUPS,
        "valid_total": valid_total 
    }

def create_analysis_report(all_analysis, bloom_check):
    """สร้างข้อมูลสรุปและ DataFrame สำหรับการแสดงผล"""
    total_questions = len(all_analysis)
    
    failed_analysis = sum(1 for a in all_analysis if "QUOTA EXCEEDED" in a.get('improvement_suggestion', '') or "เกิดข้อผิดพลาดในการวิเคราะห์" in a.get('improvement_suggestion', ''))
    
    successfully_analyzed_questions = total_questions - failed_analysis
    good_questions = sum(1 for a in all_analysis if a.get('is_good_question') is True and a.get('bloom_level') != "ไม่สามารถระบุได้")

    summary_data = {
        "สถิติโดยรวม": {
            "จำนวนข้อสอบทั้งหมด": f"{total_questions} ข้อ",
            "ข้อสอบที่วิเคราะห์สำเร็จ": f"{successfully_analyzed_questions} ข้อ",
            "ข้อสอบ **ดี** (ใช้ได้เลย)": f"{good_questions} ข้อ ({round((good_questions/successfully_analyzed_questions)*100, 1) if successfully_analyzed_questions > 0 else 0}%)",
            "ข้อสอบ **ต้องปรับปรุง**": f"{successfully_analyzed_questions - good_questions} ข้อ ({round(((successfully_analyzed_questions - good_questions)/successfully_analyzed_questions)*100, 1) if successfully_analyzed_questions > 0 else 0}%)"
        },
        "การกระจายระดับความคิด": {
            "ผลลัพธ์โดยรวม": bloom_check['reason'],
            "ระดับความคิดต่ำ (จำ/เข้าใจ) (เป้าหมาย ≤ 40%)": f"{bloom_check['percentages'].get('Remember/Understand', 0)}%",
            "ระดับความคิดกลาง (ใช้/วิเคราะห์) (เป้าหมาย ≥ 50%)": f"{bloom_check['percentages'].get('Apply/Analyze', 0)}%", 
            "ระดับความคิดสูง (ประเมิน/สร้างสรรค์) (เป้าหมาย ≥ 10%)": f"{bloom_check['percentages'].get('Evaluate/Create', 0)}%",
            "ข้อที่ระบุระดับความคิดไม่ได้": f"{bloom_check['raw_counts'].get('Unknown', 0)} ข้อ",
            "raw_counts": bloom_check['raw_counts'],
            "valid_total": bloom_check.get('valid_total', 0)
        }
    }
    
    df_data = []
    for i, item in enumerate(all_analysis):
        is_good = "&#9989; ดี" if item.get('is_good_question') is True and item.get('bloom_level') != "ไม่สามารถระบุได้" else "&#10060; ปรับปรุง/ล้มเหลว"
        df_data.append({
            'ข้อที่': i + 1,
            'คุณภาพข้อสอบ': is_good,
            'ระดับความคิด': item.get('bloom_level', 'ไม่ระบุ'),
            'มาตรฐานหลักสูตร': item.get('curriculum_standard', 'ไม่ระบุ'),
            'คำตอบ': item.get('correct_option', 'ไม่ระบุ'),
            'เหตุผลโดยย่อ': item.get('reasoning', 'ไม่ระบุเหตุผล'),
            'ข้อเสนอแนะ': item.get('improvement_suggestion', 'ไม่มี'),
            'คำถามเต็ม': item.get('question_text', '')
        })
    df = pd.DataFrame(df_data)
    return summary_data, df


def toggle_language():
    if st.session_state.language == 'th':
        st.session_state.language = 'en'
    else:
        st.session_state.language = 'th'

# --- 4. Main App Function (UI) ---

def run_app():
    # 🎨 ตั้งค่าหน้าจอ
    st.set_page_config(
        page_title="เครื่องมือวิเคราะห์ข้อสอบอัตโนมัติ (Gemini AI)",
        page_icon="📝", 
        layout="wide",
        initial_sidebar_state="auto", 
        menu_items=None
    )
    
    # 🎨 Shadcn/Tailwind CSS
    st.markdown(SHADCN_CSS, unsafe_allow_html=True)
    
    # Language toggle function

    
    # Modern Minimal Header (Dynamic)
    render_top_navigation()
    st.markdown('---')
    st.markdown(f"""
    <div style="text-align: center; padding: 1.5rem 1rem 2rem 1rem; margin-bottom: 0.5rem;">
        <h1 style="font-size: 2.4rem; font-weight: 700; background: linear-gradient(135deg, #6366f1 0%, #8b5cf6 50%, #a855f7 100%); -webkit-background-clip: text; -webkit-text-fill-color: transparent; background-clip: text; margin-bottom: 0.75rem;">
            {t('app_title')}
        </h1>
        <p style="font-size: 1.1rem; color: #4b5563; max-width: 600px; margin: 0 auto; line-height: 1.6;">
            {t('app_subtitle')}
        </p>
    </div>
    """, unsafe_allow_html=True)
    
    st.markdown("<hr>", unsafe_allow_html=True) 

    with st.sidebar:
        render_history_sidebar_v2()




    # --- Step 1: Upload ---
    
    # --- Custom Prompt (Main) ---
    st.markdown("---")
    with st.expander(t('custom_prompt_title'), expanded=False):
        st.markdown(f"**{t('custom_prompt_label')}**")
        custom_prompt_input = st.text_area(
            "Custom Prompt",
            value=st.session_state.custom_prompt,
            height=150,
            placeholder=t('custom_prompt_placeholder'),
            key='custom_prompt_main',
            label_visibility="collapsed"
        )
        if custom_prompt_input != st.session_state.custom_prompt:
            st.session_state.custom_prompt = custom_prompt_input
            st.session_state.analysis_results = None
        
        col_status1, col_status2 = st.columns([1, 1])
        with col_status1:
            if st.session_state.custom_prompt.strip():
                st.success(t('custom_prompt_active'))
            else:
                st.info(t('custom_prompt_default'))

    st.markdown("---")
    st.header(t('step1_title'))
    with st.container(border=True):
        st.markdown(f"**{t('file_uploader_label')}**")
        uploaded_file = st.file_uploader(
            t('file_uploader_label'), 
            type=['pdf', 'txt', 'docx'], 
            accept_multiple_files=False, 
            key='file_uploader_widget', 
            label_visibility="collapsed"
        )
        st.caption(t('file_tip'))

        if uploaded_file is not None:
            # ตรวจสอบว่าไฟล์เปลี่ยนไปหรือไม่
            if uploaded_file.name != st.session_state.last_uploaded_file_name:
                st.session_state.analysis_results = None
                st.session_state.last_uploaded_file_name = uploaded_file.name
                st.session_state.question_texts = None
                
            if st.session_state.question_texts is None:
                file_extension = uploaded_file.name.split('.')[-1].lower()
                
                # Custom Loading UI สำหรับการสกัดข้อสอบ
                status_container = st.empty()
                status_container.info(f"{t('reading_file')}\n\n{t('from_file')} **{uploaded_file.name}**...")
                
                with st.spinner(t('extracting')):
                    try:
                        if file_extension == 'pdf':
                            with io.BytesIO(uploaded_file.getvalue()) as open_pdf_file:
                                pdf_reader = PdfReader(open_pdf_file)
                                raw_text = extract_text_from_pdf(pdf_reader)
                        elif file_extension == 'docx':
                            raw_text = extract_text_from_docx(uploaded_file)
                            if raw_text is None:
                                raise ValueError("ไม่สามารถอ่านไฟล์ DOCX ได้ หรือไม่ได้ติดตั้ง python-docx")
                        elif file_extension == 'txt':
                            raw_text = uploaded_file.getvalue().decode("utf-8")
                        
                        question_texts = extract_questions(raw_text)
                        st.session_state.question_texts = question_texts
                        
                        status_container.empty() # ล้าง Custom Loading
                        
                        if not question_texts:
                            st.error(t('no_questions_found'))
                            st.info(t('file_tip'))
                            return 
                        
                    except Exception as e:
                        status_container.empty() # ล้าง Custom Loading
                        st.error(f"{t('file_read_error')} {e}")
                        return 
                
                # Rerun ครั้งเดียวหลังจากสกัดเสร็จ เพื่อแสดงผลลัพธ์
                st.rerun() 

            question_texts = st.session_state.question_texts
            if question_texts:
                st.success(t('extracted_questions').format(count=len(question_texts), filename=uploaded_file.name))
            


    # --- Step 2: Start Analysis ---
    st.markdown("---")
    st.header(t('step2_title'))
    
    # ใช้ Callback function เพื่อวิเคราะห์และบันทึกผลลัพธ์ (แก้ปัญหา Rerun ซ้ำซ้อน)
    def start_analysis_callback():
        # ตรวจสอบ API ตาม provider ที่เลือก
        provider = st.session_state.get('selected_provider', DEFAULT_PROVIDER)
        provider_available = (
            (provider == "Gemini (Google)" and GEMINI_AVAILABLE) or
            (provider == "Groq (ฟรี+เร็วมาก)" and GROQ_AVAILABLE) or
            (provider == "OpenRouter (หลายโมเดลฟรี)" and OPENROUTER_AVAILABLE)
        )
        
        if not provider_available:
            st.error(f"&#10060; ไม่พบ API Key สำหรับ {provider} - กรุณาตั้งค่าใน .env")
            return
            
        question_texts = st.session_state.question_texts
        analysis_results = []
        
        # ใช้ st.status เพื่อรวมสถานะทั้งหมด
        with st.status(t('starting_analysis'), expanded=True) as status_box:
            
            # ดึงโมเดลที่ผู้ใช้เลือก
            provider_models = AI_PROVIDERS.get(provider, {}).get("models", {})
            current_model = provider_models.get(st.session_state.selected_model, "unknown")
            st.write(f"⏳ กำลังวิเคราะห์ {len(question_texts)} ข้อ ด้วย `{provider}` > `{current_model}`")
            progress_bar = st.progress(0, text=t('analysis_progress').format(current=0, total=len(question_texts)))
            
            for i, q_text in enumerate(question_texts):
                st.write(t('analyzing_question').format(num=i+1))
                analysis = analyze_question(q_text, question_id=i+1)  # ใช้ wrapper function
                if "**เกิดข้อผิดพลาด" in analysis.get('improvement_suggestion', ''):
                     st.error(f"Error analyzing question {i+1}: {analysis.get('improvement_suggestion')}")

                analysis["question_text"] = q_text 
                analysis_results.append(analysis)
                
                progress_percent = (i + 1) / len(question_texts)
                progress_bar.progress(progress_percent, text=t('analysis_progress').format(current=i+1, total=len(question_texts)))
                
                # เพิ่ม delay ระหว่าง requests เพื่อหลีกเลี่ยง rate limit
                if i < len(question_texts) - 1:  # ไม่ต้องรอหลังข้อสุดท้าย
                    time.sleep(2)  # รอ 2 วินาทีระหว่างแต่ละ request
            
            # บันทึกผลลัพธ์ลงใน session state
            st.session_state.analysis_results = analysis_results
            
            # อัพเดทสถานะ
            status_box.update(label=t('analysis_complete'), state="complete", expanded=False)


    if st.session_state.question_texts and st.button(
        t('start_analysis_btn'), 
        type="primary", 
        use_container_width=True,
        on_click=start_analysis_callback 
    ):
        pass


    # --- Step 3: Report ---
    if st.session_state.analysis_results:
        st.divider()
        st.header(t('step3_title'))

        all_analysis = st.session_state.analysis_results
        successful_analysis = [a for a in all_analysis if a.get('bloom_level') != "ไม่สามารถระบุได้"]
        bloom_check = check_bloom_criteria(successful_analysis)
        summary_data, df = create_analysis_report(all_analysis, bloom_check)
        
        # ดึง valid_total ออกมาเพื่อใช้ในการคำนวณสัดส่วนในตาราง (แก้ NameError)
        valid_total = summary_data["การกระจายระดับความคิด"].get("valid_total", 0)
        
        # บันทึกประวัติการวิเคราะห์
        save_analysis_history(
            st.session_state.get('last_uploaded_file_name', 'unknown'),
            all_analysis,
            summary_data
        )

        # ใช้ Tabs สำหรับจัดระเบียบรายงาน - เพิ่ม Export tab
        tab_summary, tab_details, tab_export = st.tabs([
            t('tab_summary'), 
            t('tab_details'),
            "&#128229; Export รายงาน"
        ])

        # --- Tab: Summary ---
        with tab_summary:
            st.subheader(t('summary_title'))
            stats = summary_data["สถิติโดยรวม"]
            col1, col2, col3, col4 = st.columns(4) 

            # Helper function to extract percent value
            def get_percent_delta(text):
                try: 
                    return text.split('(')[1].strip('%)') 
                except IndexError: 
                    return "0.0%"

            # Good Questions Metric
            good_count_str = stats["ข้อสอบ **ดี** (ใช้ได้เลย)"].split(' ')[0]
            good_percent_str = get_percent_delta(stats["ข้อสอบ **ดี** (ใช้ได้เลย)"])
            col1.metric(t('good_questions'), good_count_str, delta=f"{good_percent_str}%", delta_color="normal")
            
            # To Improve Metric
            improve_count_str = stats["ข้อสอบ **ต้องปรับปรุง**"].split(' ')[0]
            improve_percent_str = get_percent_delta(stats["ข้อสอบ **ต้องปรับปรุง**"])
            col2.metric(t('needs_improvement'), improve_count_str, delta=f"{improve_percent_str}%", delta_color="inverse")
            
            # Total Questions 
            col3.metric(t('total_questions'), stats["จำนวนข้อสอบทั้งหมด"].split(' ')[0])
            
            # Successfully Analyzed
            col4.metric(t('analyzed_success'), stats["ข้อสอบที่วิเคราะห์สำเร็จ"].split(' ')[0])
            
            st.markdown("---")
            st.subheader(t('bloom_criteria_title'))
            bloom_stats = summary_data["การกระจายระดับความคิด"]
            
            if bloom_check['pass']:
                st.success(f"**🎉 {bloom_stats['ผลลัพธ์โดยรวม']}**")
            else:
                st.warning(f"**&#10060; {bloom_stats['ผลลัพธ์โดยรวม']}**")
                
            col_b1, col_b2, col_b3 = st.columns(3)
            col_b1.metric(t('bloom_low'), bloom_stats["ระดับความคิดต่ำ (จำ/เข้าใจ) (เป้าหมาย ≤ 40%)"], delta=f"{t('target')} ≤ 40%")
            col_b2.metric(t('bloom_mid'), bloom_stats["ระดับความคิดกลาง (ใช้/วิเคราะห์) (เป้าหมาย ≥ 50%)"], delta=f"{t('target')} ≥ 50%")
            col_b3.metric(t('bloom_high'), bloom_stats["ระดับความคิดสูง (ประเมิน/สร้างสรรค์) (เป้าหมาย ≥ 10%)"], delta=f"{t('target')} ≥ 10%")
            
            st.markdown(f"{t('unidentified_bloom')} {bloom_stats['ข้อที่ระบุระดับความคิดไม่ได้']}")
            
            
            # --- สร้าง Pie Chart และ ตารางสรุป ---
            st.markdown("---")
            st.subheader(t('bloom_distribution'))
            
            col_chart, col_table = st.columns([1, 1.2]) 

            # 1. Pie Chart 
            with col_chart:
                bloom_counts = bloom_stats['raw_counts']
                # เตรียมข้อมูลสำหรับ Pie Chart (ไม่รวม Unknown)
                chart_data_raw = {
                    'ระดับ Bloom': list(bloom_counts.keys())[:-1],
                    'จำนวนข้อ': list(bloom_counts.values())[:-1],
                    'สี': [get_bloom_color(level) for level in list(bloom_counts.keys())[:-1]]
                }
                chart_df = pd.DataFrame(chart_data_raw)
                
                if not chart_df.empty and chart_df['จำนวนข้อ'].sum() > 0:
                    base = alt.Chart(chart_df).encode(
                        theta=alt.Theta("จำนวนข้อ", stack=True)
                    )
                    
                    pie = base.mark_arc(outerRadius=120).encode(
                        color=alt.Color("ระดับ Bloom", scale=alt.Scale(domain=chart_df['ระดับ Bloom'].tolist(), range=chart_df['สี'].tolist())),
                        order=alt.Order("จำนวนข้อ", sort="descending"),
                        tooltip=["ระดับ Bloom", "จำนวนข้อ"]
                    )
                    
                    st.altair_chart(pie, use_container_width=True)
                else:
                    st.warning("ไม่มีข้อมูลข้อสอบที่วิเคราะห์ระดับ Bloom ได้")
            
            # 2. ตารางสรุปคุณภาพ/จำนวนข้อตามระดับ Bloom
            with col_table:
                # สร้าง DataFrame สรุป
                summary_table_data = []
                for level, count in bloom_stats['raw_counts'].items():
                    if level == 'Unknown': continue 
                    
                    level_items = [a for a in all_analysis if level.lower() in a.get('bloom_level', '').lower()]
                    good_count = sum(1 for a in level_items if a.get('is_good_question') is True)
                    
                    percent_text = f"{round((count / valid_total) * 100, 1) if valid_total > 0 else 0}%"
                    
                    summary_table_data.append({
                        'ระดับ Bloom': level,
                        'จำนวนข้อ': count,
                        'สัดส่วน': percent_text,
                        'คุณภาพดี': good_count,
                        'ต้องปรับปรุง': count - good_count
                    })
                
                summary_table_df = pd.DataFrame(summary_table_data)
                
                st.markdown("**ตารางสรุปจำนวนและคุณภาพข้อสอบตามระดับ Bloom**")
                st.dataframe(
                    summary_table_df, 
                    hide_index=True, 
                    use_container_width=True,
                    column_config={
                        'จำนวนข้อ': st.column_config.NumberColumn(format="%d ข้อ"),
                        'คุณภาพดี': st.column_config.NumberColumn(format="%d ข้อ"),
                        'ต้องปรับปรุง': st.column_config.NumberColumn(format="%d ข้อ"),
                    }
                )

        # --- Tab: Details ---
        with tab_details:
            st.subheader(t('details_title'))
            
            # 1. แสดง DataFrame สรุปก่อน
            st.dataframe(
                df[[
                    'ข้อที่', 'คุณภาพข้อสอบ', 'ระดับความคิด', 
                    'มาตรฐานหลักสูตร', 'คำตอบ', 'เหตุผลโดยย่อ', 
                    'ข้อเสนอแนะ'
                ]],
                column_config={
                    "คุณภาพข้อสอบ": st.column_config.Column("คุณภาพข้อสอบ", width="small"),
                    "ระดับความคิด": st.column_config.Column("ระดับความคิด", width="small"),
                    "เหตุผลโดยย่อ": st.column_config.Column("เหตุผลโดยย่อ", width="medium"),
                },
                use_container_width=True,
                hide_index=True
            )
            
            # 2. Loop สร้าง expander สำหรับทุกข้อ
            st.markdown("---")
            st.markdown(t('click_detail'))
            
            for q_index, item in enumerate(all_analysis):
                quality_status = t('good') if item.get('is_good_question') is True and item.get('bloom_level') != "ไม่สามารถระบุได้" else t('improve')
                expander_title = f"**{t('question_num')} {q_index+1}** | {quality_status} | {t('bloom_level')}: **{item.get('bloom_level', 'ไม่ระบุ')}**"
                
                # ใช้ st.expander เพื่อแสดงรายละเอียด
                with st.expander(expander_title):
                    
                    st.markdown(t('full_question'))
                    st.code(item.get('question_text', 'N/A'), language='markdown')
                    
                    st.markdown("---")

                    # การแสดงผลแบบชิดและมีสีสันสวยงาม 
                    bloom_color = get_bloom_color(item.get('bloom_level', 'ไม่ระบุ'))
                    text_color = get_text_color_for_bloom(item.get('bloom_level', 'ไม่ระบุ'))
                    
                    col_det1, col_det2, col_det3 = st.columns([1, 1, 1]) 

                    with col_det1:
                        # แสดง Bloom Level
                        st.markdown(
                            f"""
                            <div style='background-color:{bloom_color}; color:{text_color}; padding: 10px; border-radius: 5px; text-align: center; margin-bottom: 10px;'>
                                <strong>&#128161; ระดับ Bloom:</strong> {item.get('bloom_level', 'N/A')}
                            </div>
                            """, unsafe_allow_html=True
                        )
                    
                    with col_det2:
                         # แสดง Difficulty (พร้อมแก้ไขสีตัวอักษร)
                        difficulty_level = item.get('difficulty', 'ไม่ระบุ')
                        difficulty_color = {"ง่าย": "#008000", "ปานกลาง": "#FFA500", "ยาก": "#FF4500"}.get(difficulty_level, "#808080")
                        
                        # กำหนดสีตัวอักษรตามพื้นหลัง
                        if difficulty_level in ["ปานกลาง"]: 
                            difficulty_text_color = "white" 
                        else:
                            difficulty_text_color = "white"
                        
                        st.markdown(
                            f"""
                            <div style='background-color:{difficulty_color}; color:{difficulty_text_color}; padding: 10px; border-radius: 5px; text-align: center; margin-bottom: 10px;'>
                                <strong>{t('difficulty')}</strong> {difficulty_level}
                            </div>
                            """, unsafe_allow_html=True
                        )

                    with col_det3:
                        # แสดง Correct Option
                        st.markdown(
                            f"""
                            <div style='background-color:#0077B6; color:white; padding: 10px; border-radius: 5px; text-align: center; margin-bottom: 10px;'>
                                <strong>{t('correct_answer')}</strong> {item.get('correct_option', 'N/A')}
                            </div>
                            """, unsafe_allow_html=True
                        )

                    # ข้อมูลหลัก 
                    st.markdown(f"{t('curriculum_indicator')} `{item.get('curriculum_standard', 'N/A')}`")
                    st.markdown(f"{t('bloom_reason')} {item.get('reasoning', 'N/A')}")
                    
                    st.divider()
                    st.subheader(t('answer_analysis_title'))
                    st.markdown(f"{t('correct_analysis')} {item.get('correct_option_analysis', 'N/A')}")
                    st.markdown(f"{t('distractor_analysis')} {item.get('distractor_analysis', 'N/A')}")
                    st.markdown(f"{t('why_good_distractor')} {item.get('why_good_distractor', 'N/A')}")
                    st.warning(f"{t('improvement_suggestion')} {item.get('improvement_suggestion', 'N/A')}")
                    
                    # ปุ่มปรับปรุงข้อสอบด้วย AI
                    if st.button(f"&#10024; ปรับปรุงข้อสอบข้อที่ {q_index+1}", key=f"improve_{q_index}"):
                        with st.spinner("กำลังปรับปรุงข้อสอบ..."):
                            improved, err = improve_question_with_ai(
                                item.get('question_text', ''),
                                item.get('improvement_suggestion', '')
                            )
                            if improved:
                                st.success("&#9989; ข้อสอบที่ปรับปรุงแล้ว:")
                                st.markdown(improved)
                            else:
                                st.error(f"&#10060; เกิดข้อผิดพลาด: {err}")
                
                st.divider() 


        # --- Tab: Export ---
        with tab_export:
            st.subheader("&#128229; ดาวน์โหลดรายงานผลวิเคราะห์")
            
            col_excel, col_info = st.columns([1, 2])
            
            with col_excel:
                if EXCEL_AVAILABLE:
                    excel_data = export_to_excel(all_analysis)
                    if excel_data:
                        st.download_button(
                            label="📊 ดาวน์โหลด Excel",
                            data=excel_data,
                            file_name=f"exam_analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            use_container_width=True
                        )
                else:
                    st.warning("ต้องติดตั้ง openpyxl เพื่อใช้ฟีเจอร์นี้")
            
            with col_info:
                st.info(f"""
                **ข้อมูลที่จะ Export:**
                - จำนวนข้อสอบ: {len(all_analysis)} ข้อ
                - ข้อสอบคุณภาพดี: {sum(1 for a in all_analysis if a.get('is_good_question'))} ข้อ
                - ข้อสอบต้องปรับปรุง: {len(all_analysis) - sum(1 for a in all_analysis if a.get('is_good_question'))} ข้อ
                """)

    # --- Section: Generate New Exam ---
    st.divider()
    st.header("&#127381; สร้างข้อสอบใหม่ด้วย AI")
    
    with st.expander("&#128161; คลิกเพื่อสร้างข้อสอบใหม่", expanded=False):
        col_gen1, col_gen2 = st.columns(2)
        
        with col_gen1:
            subject = st.text_input("&#128218; วิชา/หัวข้อ", placeholder="เช่น คณิตศาสตร์ ม.3, วิทยาศาสตร์ ป.6")
            bloom_level = st.selectbox(
                "&#129504; ระดับ Bloom's Taxonomy",
                ["จำ (Remember)", "เข้าใจ (Understand)", "ประยุกต์ใช้ (Apply)", 
                 "วิเคราะห์ (Analyze)", "ประเมินค่า (Evaluate)", "สร้างสรรค์ (Create)"]
            )
        
        with col_gen2:
            num_questions = st.number_input("📝 จำนวนข้อ", min_value=1, max_value=20, value=5)
            difficulty = st.selectbox("📊 ระดับความยาก", ["ง่าย", "ปานกลาง", "ยาก"])
        
        if st.button("&#128640; สร้างข้อสอบ", type="primary", use_container_width=True):
            if not subject:
                st.warning("กรุณาระบุวิชา/หัวข้อ")
            else:
                with st.spinner(f"กำลังสร้างข้อสอบ {num_questions} ข้อ..."):
                    exams, err = generate_exam_with_ai(subject, bloom_level, num_questions, difficulty)
                    if exams:
                        st.success(f"&#9989; สร้างข้อสอบสำเร็จ {len(exams)} ข้อ!")
                        for i, exam in enumerate(exams, 1):
                            with st.expander(f"ข้อ {i}: {exam.get('question', 'N/A')[:50]}..."):
                                st.markdown(f"**คำถาม:** {exam.get('question', 'N/A')}")
                                st.markdown("**ตัวเลือก:**")
                                options = exam.get('options', [])
                                for j, opt in enumerate(options):
                                    prefix = ['ก.', 'ข.', 'ค.', 'ง.'][j] if j < 4 else f"{j+1}."
                                    st.markdown(f"   {prefix} {opt}")
                                st.markdown(f"**เฉลย:** {exam.get('answer', 'N/A')}")
                                st.markdown(f"**คำอธิบาย:** {exam.get('explanation', 'N/A')}")
                    else:
                        st.error(f"&#10060; เกิดข้อผิดพลาด: {err}")

    # --- Section: History ---
    st.divider()
    st.header("&#128220; ประวัติการวิเคราะห์")
    
    history = load_analysis_history()
    if history:
        for entry in history[:5]:  # แสดง 5 รายการล่าสุด
            timestamp = entry.get('timestamp', 'N/A')[:10]
            filename = entry.get('filename', 'N/A')
            total = entry.get('total_questions', 0)
            good = entry.get('good_questions', 0)
            st.markdown(f"📁 **{filename}** - {timestamp} | {total} ข้อ (ดี: {good})")
    else:
        st.info("ยังไม่มีประวัติการวิเคราะห์")


if __name__ == "__main__":
    run_app()