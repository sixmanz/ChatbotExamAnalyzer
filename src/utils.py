# -*- coding: utf-8 -*-
import os
import re
import json
import io
import streamlit as st
import pandas as pd
from datetime import datetime
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment

# --- Optional Imports ---
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

# --- Colors ---
BLOOM_COLORS = {
    "REMEMBER": "#BFDBFE",   # Blue-200
    "UNDERSTAND": "#A5F3FC", # Cyan-200
    "APPLY": "#BBF7D0",      # Green-200
    "ANALYZE": "#FEF08A",    # Yellow-200
    "EVALUATE": "#FED7AA",   # Orange-200
    "CREATE": "#FBCFE8",     # Pink-200
    "ไม่ระบุ": "#F1F5F9",    # Slate-100
    "ไม่สามารถระบุได้": "#F1F5F9"
}

def get_bloom_color(level):
    # Returns Hex color for Bloom's Taxonomy level
    if not level:
        return BLOOM_COLORS['ไม่ระบุ']
    
    level_upper = level.strip().upper()
    for key, color in BLOOM_COLORS.items():
        if key in level_upper:
            return color
    return BLOOM_COLORS['ไม่ระบุ'] # Fallback

def get_text_color_for_bloom(level):
    """กำหนดสีตัวอักษรให้ตัดกับสีพื้นหลัง (Pastel ใช้สีเข้ม)"""
    return '#1e293b' # Slate-800 for all pastel backgrounds

# --- File Handling ---
def extract_text_from_pdf(file):
    """สกัดข้อความจากไฟล์ PDF (Cached)"""
    from PyPDF2 import PdfReader
    text = ""
    try:
        pdf_reader = PdfReader(file)
        for page in pdf_reader.pages:
            try:
                page_text = page.get_text() if hasattr(page, 'get_text') else page.extract_text()
                text += (page_text or "") + "\n"
            except Exception:
                text += "\n"
    except Exception:
        pass
    return text.strip()

def extract_text_from_docx(file):
    """สกัดข้อความจากไฟล์ DOCX (อ่านพารากราฟและตาราง)"""
    if not DOCX_AVAILABLE:
        return None
    try:
        doc = Document(file)
        full_text = []
        
        # อ่าน paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
                
        # อ่าน tables (มักใช้ในข้อสอบ)
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    full_text.append(" ".join(row_text))
                    
        return "\n".join(full_text)
    except Exception as e:
        return None

def clean_and_normalize(text):
    """ทำความสะอาดข้อความและแปลงเลขไทยเป็นเลขอารบิก"""
    if not text: return ""
    # 1. แปลงเลขไทย
    thai_digits = "๐๑๒๓๔๕๖๗๘๙"
    for i, digit in enumerate(thai_digits):
        text = text.replace(digit, str(i))
    
    # 2. ทำความสะอาดตัวอักษรพิเศษและช่องว่าง
    text = re.sub(r'[ \t]+', ' ', text) 
    text = text.replace('\r', '')
    
    # 3. ทำให้ตัวเลือกติดกับจุด (เช่น ก . -> ก.)
    text = re.sub(r'([ก-งA-D])\s*\.', r'\1.', text)
    
    # 4. ทำให้เลขข้อติดกับจุด (เช่น 1 . -> 1.)
    text = re.sub(r'(\d+)\s*\.', r'\1.', text)
    
    # NEW: Insert newline before potential question start if preceded by punctuation or space
    # Matches: "text. 2. text" -> "text.\n2. text"
    # Matches: "text (2) text" -> "text\n(2) text"
    text = re.sub(r'(\s+)(\(?\d+[\.\)])\s', r'\n\2 ', text)
    
    # 5. ลบบรรทัดว่างที่ติดกันหลายบรรทัด
    text = re.sub(r'\n{2,}', '\n', text)
    
    lines = [line.strip() for line in text.split('\n')]
    return '\n'.join(lines)

# --- Prompts ---
def load_prompts(prompt_path_override=None) -> tuple[str, str]:
    """อ่านเนื้อหา Prompt Template จากไฟล์ 'Prompt.txt' และแยกเป็น System/Chat"""
    if prompt_path_override:
        prompt_path = prompt_path_override
    else:
        # Default assumption: Prompt.txt is in the same dir as this file or one level up
        # Try both current dir and parent dir
        current_dir = os.path.dirname(os.path.abspath(__file__))
        prompt_path = os.path.join(current_dir, 'Prompt.txt')
        if not os.path.exists(prompt_path):
             prompt_path = os.path.join(current_dir, '..', 'Prompt.txt')

    try:
        with open(prompt_path, 'r', encoding='utf-8') as f:
            full_content = f.read()
            
            # 1. System Instruction
            system_match = full_content.split("# --- END_SYSTEM_INSTRUCTION_ANALYSIS_MODE ---")[0]
            SYSTEM_INSTRUCTION_PROMPT = system_match.strip() if system_match else "You are a test expert for Thai curriculum. Analyze the question and return a JSON object."
            
            # 2. Few-Shot Template 
            chat_match = full_content.split("# --- CHAT_PROMPT ---")
            
            if len(chat_match) > 1:
                # แยกส่วน Prompt Template ทั้งหมด
                template_content = chat_match[1].split("# --- CHAT_PROMPT_END ---")[0].strip()
                # และเพิ่ม {user_query} เข้าไปในส่วนท้าย
                FEW_SHOT_PROMPT_TEMPLATE = template_content + "\n{user_query}"
            else:
                FEW_SHOT_PROMPT_TEMPLATE = "Analyze the following question/text and return the JSON object: {user_query}"

            return SYSTEM_INSTRUCTION_PROMPT, FEW_SHOT_PROMPT_TEMPLATE
    except FileNotFoundError:
        return (
            "# Error: Prompt file not found. Fallback to default instruction.", 
            "Analyze the following question/text and return the JSON object: {user_query}"
        )

# --- Analysis Helpers ---
def create_error_response(error_message):
    """สร้าง response เมื่อเกิดข้อผิดพลาด"""
    return {
        "bloom_level": "ไม่สามารถระบุได้", "reasoning": "AI วิเคราะห์ล้มเหลว",
        "difficulty": "ไม่สามารถประเมินได้", "curriculum_standard": "ไม่สามารถระบุได้",
        "correct_option": "ไม่ระบุ", "correct_option_analysis": "ไม่ระบุ",
        "distractor_analysis": "ไม่ระบุ", "why_good_distractor": "ไม่ระบุ",
        "is_good_question": False, 
        "improvement_suggestion": f"**เกิดข้อผิดพลาด**: {error_message}"
    }

def sanitize_analysis(analysis):
    """ทำความสะอาดและตรวจสอบผลลัพธ์จาก AI (Robust)"""
    required_keys = [
        "bloom_level", "reasoning", "difficulty", "curriculum_standard",
        "correct_option", "correct_option_analysis", "distractor_analysis",
        "why_good_distractor", "is_good_question", "improvement_suggestion"
    ]
    
    result = {}
    missing_keys = []
    
    for key in required_keys:
        val = analysis.get(key)
        
        # Robust Boolean Conversion
        if key == "is_good_question":
            if isinstance(val, bool):
                result[key] = val
            elif isinstance(val, str):
                result[key] = val.strip().lower() in ['true', 'yes', '1', 'correct', 'จริง', 'ใช่']
            else:
                result[key] = False
        else:
            # Robust String Conversion
            result[key] = str(val).strip() if val not in [None, "", "null"] else "ไม่ระบุ"
            
        if key not in analysis:
            missing_keys.append(key)

    # Specific Default Values for missing keys
    if "improvement_suggestion" not in result or result["improvement_suggestion"] == "ไม่ระบุ":
        result["improvement_suggestion"] = "ไม่มีข้อเสนอแนะเพิ่มเติม"

    return result

def check_bloom_criteria(analysis_results):
    """ตรวจสอบว่าชุดข้อสอบผ่านเกณฑ์การกระจายระดับ Bloom หรือไม่."""
    total = len(analysis_results)
    if total == 0: 
        return {"pass": False, "reason": "ไม่มีข้อมูลข้อสอบ", "percentages": {}, "raw_counts": {}, "valid_total": 0} 

    # Standardized Level Names for grouping
    LEVEL_GROUPS = {
        "Remember": 0, "Understand": 0, "Apply": 0, 
        "Analyze": 0, "Evaluate": 0, "Create": 0, "Unknown": 0
    }
    
    for item in analysis_results:
        level = str(item.get("bloom_level", "")).strip().lower()
        if "remember" in level: LEVEL_GROUPS["Remember"] += 1
        elif "understand" in level: LEVEL_GROUPS["Understand"] += 1
        elif "apply" in level: LEVEL_GROUPS["Apply"] += 1
        elif "analyze" in level: LEVEL_GROUPS["Analyze"] += 1
        elif "evaluate" in level: LEVEL_GROUPS["Evaluate"] += 1
        elif "create" in level: LEVEL_GROUPS["Create"] += 1
        else: LEVEL_GROUPS["Unknown"] += 1

    apply_analyze_count = LEVEL_GROUPS["Apply"] + LEVEL_GROUPS["Analyze"]
    remember_understand_count = LEVEL_GROUPS["Remember"] + LEVEL_GROUPS["Understand"]
    evaluate_create_count = LEVEL_GROUPS["Evaluate"] + LEVEL_GROUPS["Create"]
    valid_total = total - LEVEL_GROUPS["Unknown"] # ข้อที่สามารถระบุระดับ Bloom ได้

    if valid_total == 0:
        return {"pass": False, "reason": "ไม่มีข้อสอบที่สามารถวิเคราะห์ระดับความคิดได้เลย", "percentages": {}, "raw_counts": LEVEL_GROUPS, "valid_total": 0} 
        
    percent = {
        "Remember/Understand": round((remember_understand_count / valid_total) * 100, 1) if valid_total > 0 else 0,
        "Apply/Analyze": round((apply_analyze_count / valid_total) * 100, 1) if valid_total > 0 else 0,
        "Evaluate/Create": round((evaluate_create_count / valid_total) * 100, 1) if valid_total > 0 else 0
    }

    # เกณฑ์ (ตัวอย่าง: ต่ำ ≤ 40%, กลาง ≥ 50%, สูง ≥ 10%)
    pass_r_u = percent["Remember/Understand"] <= 40
    pass_a_a = percent["Apply/Analyze"] >= 50
    pass_e_c = percent["Evaluate/Create"] >= 10
    passed = pass_r_u and pass_a_a and pass_e_c

    reason = "ชุดข้อสอบ **ผ่าน** เกณฑ์การกระจายระดับความคิด (Bloom)" if passed else "ชุดข้อสอบ **ไม่ผ่าน** เกณฑ์การกระจายระดับความคิด (Bloom)"

    return {
        "pass": passed,
        "reason": reason,
        "percentages": percent,
        "raw_counts": LEVEL_GROUPS,
        "valid_total": valid_total 
    }

# --- Export/History ---
def export_to_excel(analysis_results, filename="exam_analysis.xlsx"):
    """Export ผลวิเคราะห์เป็น Excel"""
    if not EXCEL_AVAILABLE:
        return None
    
    wb = Workbook()
    ws = wb.active
    ws.title = "ผลวิเคราะห์ข้อสอบ"
    
    # Header styling
    header_fill = PatternFill(start_color="18181B", end_color="18181B", fill_type="solid")
    header_font = Font(color="FFFFFF", bold=True)
    
    headers = ["ข้อที่", "ระดับ Bloom", "ความยาก", "คุณภาพ", "มาตรฐาน", "คำตอบ", "ข้อเสนอแนะ"]
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center")
    
    # Data rows
    for idx, item in enumerate(analysis_results, 1):
        ws.cell(row=idx+1, column=1, value=idx)
        ws.cell(row=idx+1, column=2, value=item.get('bloom_level', 'N/A'))
        ws.cell(row=idx+1, column=3, value=item.get('difficulty', 'N/A'))
        ws.cell(row=idx+1, column=4, value="ดี" if item.get('is_good_question') else "ต้องปรับปรุง")
        ws.cell(row=idx+1, column=5, value=item.get('curriculum_standard', 'N/A'))
        ws.cell(row=idx+1, column=6, value=item.get('correct_option', 'N/A'))
        ws.cell(row=idx+1, column=7, value=item.get('improvement_suggestion', 'N/A'))
    
    # Auto-adjust column widths
    for col in ws.columns:
        max_length = 0
        for cell in col:
             val = str(cell.value or "")
             max_length = max(max_length, len(val))
        ws.column_dimensions[col[0].column_letter].width = min(max_length + 2, 50)
    
    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    return output

def export_to_word(analysis_results, filename="exam_analysis.docx"):
    """Export ผลวิเคราะห์เป็น MS Word (.docx)"""
    if not DOCX_AVAILABLE:
        return None
        
    doc = Document()
    doc.add_heading('รายงานการวิเคราะห์คุณภาพข้อสอบ', 0)
    
    doc.add_paragraph(f"สร้างเมื่อ: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    
    # Summary
    total = len(analysis_results)
    good = sum(1 for r in analysis_results if r.get('is_good_question'))
    doc.add_heading('สรุปภาพรวม', level=1)
    doc.add_paragraph(f"จำนวนข้อสอบทั้งหมด: {total} ข้อ")
    doc.add_paragraph(f"ข้อสอบคุณภาพดี: {good} ข้อ")
    doc.add_paragraph(f"ข้อสอบต้องปรับปรุง: {total - good} ข้อ")
    
    doc.add_heading('รายละเอียดรายข้อ', level=1)
    
    for idx, item in enumerate(analysis_results, 1):
        # Title per question
        p = doc.add_paragraph()
        runner = p.add_run(f"ข้อที่ {idx}")
        runner.bold = True
        runner.font.size = 14
        
        # Table for details
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'
        
        def add_row(label, value):
            row_cells = table.add_row().cells
            row_cells[0].text = label
            row_cells[0].paragraphs[0].runs[0].bold = True
            row_cells[1].text = str(value)
            
        add_row("ระดับ Bloom", item.get('bloom_level', '-'))
        add_row("ความยาก", item.get('difficulty', '-'))
        add_row("ผลการประเมิน", "✅ ดี" if item.get('is_good_question') else "⚠️ ต้องปรับปรุง")
        add_row("มาตรฐาน", item.get('curriculum_standard', '-'))
        add_row("ข้อเสนอแนะ", item.get('improvement_suggestion', '-'))
        
        doc.add_paragraph("") # Spacing

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output

# Wrapper for Database History (To keep API consistent)
from .database import save_exam_result, get_recent_exams, load_exam_results, clear_all_history

def save_analysis_history(filename, results, summary):
    return save_exam_result(filename, results, summary)

def load_analysis_history():
    return get_recent_exams()
